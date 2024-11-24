import docx

document = docx.Document()

#writing data into the word document in various styles and formats
document.add_paragraph('Hello Word!', 'Title')
document.add_paragraph('By Tim', 'Heading 1')
document.add_paragraph('This is a word document created in Python')
document.add_paragraph('Automate the boring stuff.', 'Quote')
document.add_paragraph('List of Favorite Colors', 'Heading 2')

#create a list to add to the document
color_favorites = ['Blue', 'Purple', 'Red']

#for loop to iterate each list item with a bullet point
for color in color_favorites:
    document.add_paragraph(color, 'List Bullet')


document.save('Hello_word.docx') #save document.