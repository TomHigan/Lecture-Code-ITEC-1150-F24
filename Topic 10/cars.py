import docx


#dictionary of electric cars
electric_cars = {
    'Chevy': 'Volt',
    'Nissan': 'Leaf',
    'Tesla': 'Model S',
    'Volkswagen': 'e-Golf'
}

document.add_paragraph('Electric Cars', 'Heading 1') #creates heading in word doc

#for loop that iterates each key and value pair in the dictionary electric_cars
for make, model in electric_cars.items():
    document.add_paragraph(make, 'Heading 3') #uses the key to make a heading in the document
    document.add_paragraph(f'An electric car by {make} is {model}.') #a sentence using the key/value pairs
    # document.add_paragraph('Make: ' + make)
    # document.add_paragraph('Model: ' + model)

document.save('electric_cars.docx') #saves document
